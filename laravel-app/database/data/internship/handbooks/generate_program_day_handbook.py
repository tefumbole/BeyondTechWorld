#!/usr/bin/env python3
"""
Generate Beyond internship Student Handbook DOCX files for any program/day.

Same Day-1 pattern for all 11 programs × 180 days:
title page, schedule, concepts, install/verify, lab, responsible checkpoint,
evaluation, troubleshooting, evidence, report, checklists, rubric, references.

Usage:
  python3 generate_program_day_handbook.py --program ARTIFICIAL_INTELLIGENCE --days 6-180
  python3 generate_program_day_handbook.py --program NETWORKING --days 1-180
  python3 generate_program_day_handbook.py --all-programs --days 1-180
"""

from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SEED = HERE.parent / "beyond_180_day_curriculum_seed.json"
AI_DAY001 = HERE / "ai" / "Beyond_AI_Internship_Day_001_Student_Handbook.docx"

BLUE = RGBColor(0x0B, 0x3F, 0x90)
DARK = RGBColor(0x1E, 0x29, 0x3B)
WARN = RGBColor(0x9A, 0x34, 0x12)

PROGRAM_META = {
    "ARTIFICIAL_INTELLIGENCE": {
        "title": "Beyond Artificial Intelligence Internship",
        "folder": "ai",
        "file_prefix": "Beyond_AI_Internship",
        "discipline": "AI / data / software lab",
        "safety": "Never paste client, patient, student, employee, financial, credential, or other private data into a public AI service. Use synthetic or approved public data only.",
        "responsible_label": "Responsible-AI checkpoint",
        "toolchain_note": "Python + venv + VS Code + Jupyter from Day 1 remain the default toolchain unless the day introduces an approved addition.",
        "official_downloads": [
            ("Python", "https://www.python.org/downloads/"),
            ("VS Code", "https://code.visualstudio.com/"),
            ("Jupyter docs", "https://docs.jupyter.org/"),
        ],
        "refs": [
            "Python docs — https://docs.python.org/3/",
            "VS Code docs — https://code.visualstudio.com/docs",
            "Jupyter docs — https://docs.jupyter.org/",
            "NIST AI RMF — https://www.nist.gov/itl/ai-risk-management-framework",
            "NIST Privacy Framework — https://www.nist.gov/privacy-framework",
            "OWASP — https://owasp.org/",
        ],
    },
    "MACHINE_LEARNING": {
        "title": "Beyond Machine Learning Internship",
        "folder": "machine_learning",
        "file_prefix": "Beyond_ML_Internship",
        "discipline": "ML / data science lab",
        "safety": "Treat datasets as sensitive unless marked public. Do not scrape personal data. Keep train/test separation honest.",
        "responsible_label": "Responsible-ML checkpoint",
        "toolchain_note": "Python, Jupyter, pandas, and scikit-learn on an authorized laptop/lab VM.",
        "official_downloads": [
            ("Python", "https://www.python.org/downloads/"),
            ("VS Code", "https://code.visualstudio.com/"),
            ("scikit-learn", "https://scikit-learn.org/stable/"),
        ],
        "refs": [
            "scikit-learn user guide — https://scikit-learn.org/stable/user_guide.html",
            "pandas docs — https://pandas.pydata.org/docs/",
            "NIST AI RMF — https://www.nist.gov/itl/ai-risk-management-framework",
        ],
    },
    "DATA_SCIENCE": {
        "title": "Beyond Data Science Internship",
        "folder": "data_science",
        "file_prefix": "Beyond_DataScience_Internship",
        "discipline": "data analysis lab",
        "safety": "Anonymize personal data. Do not invent statistics. Keep raw vs cleaned datasets separated.",
        "responsible_label": "Data ethics checkpoint",
        "toolchain_note": "Python, Jupyter, pandas, LibreOffice Calc, and approved public datasets.",
        "official_downloads": [
            ("Python", "https://www.python.org/downloads/"),
            ("LibreOffice", "https://www.libreoffice.org/download/download/"),
            ("pandas", "https://pandas.pydata.org/"),
        ],
        "refs": [
            "pandas docs — https://pandas.pydata.org/docs/",
            "NIST Privacy Framework — https://www.nist.gov/privacy-framework",
        ],
    },
    "SOFTWARE_DEVELOPMENT": {
        "title": "Beyond Software Development Internship",
        "folder": "software_development",
        "file_prefix": "Beyond_SoftwareDev_Internship",
        "discipline": "software engineering lab",
        "safety": "Never commit passwords or API keys. Test only on local/dev environments. Get approval before deploying shared systems.",
        "responsible_label": "Secure development checkpoint",
        "toolchain_note": "VS Code, Git, GitHub (or Beyond-approved Git host), local databases, open-source stacks.",
        "official_downloads": [
            ("VS Code", "https://code.visualstudio.com/"),
            ("Git", "https://git-scm.com/downloads"),
            ("Python/Node as assigned", "supervisor-approved official sites only"),
        ],
        "refs": [
            "Git documentation — https://git-scm.com/doc",
            "OWASP Top Ten — https://owasp.org/www-project-top-ten/",
            "VS Code docs — https://code.visualstudio.com/docs",
        ],
    },
    "NETWORKING": {
        "title": "Beyond Networking Internship",
        "folder": "networking",
        "file_prefix": "Beyond_Networking_Internship",
        "discipline": "networking lab",
        "safety": "Do not scan or reconfigure production networks. Use Beyond lab VLANs/VMs/Packet Tracer only. Ask before touching shared infrastructure.",
        "responsible_label": "Network safety checkpoint",
        "toolchain_note": "Packet Tracer, VirtualBox, Wireshark, and Beyond-authorized lab equipment.",
        "official_downloads": [
            ("Cisco Packet Tracer (via NetAcad/supervisor)", "supervisor-provided official channel"),
            ("VirtualBox", "https://www.virtualbox.org/"),
            ("Wireshark", "https://www.wireshark.org/download.html"),
        ],
        "refs": [
            "Wireshark docs — https://www.wireshark.org/docs/",
            "VirtualBox manual — https://www.virtualbox.org/manual/",
            "NIST Cybersecurity Framework — https://www.nist.gov/cyberframework",
        ],
    },
    "CYBER_SECURITY": {
        "title": "Beyond Cyber Security Internship",
        "folder": "cyber_security",
        "file_prefix": "Beyond_CyberSecurity_Internship",
        "discipline": "cybersecurity lab",
        "safety": "Work only in isolated labs with written authorization. Never use offensive tools against live systems, classmates, or the internet.",
        "responsible_label": "Ethics & authorization checkpoint",
        "toolchain_note": "VirtualBox, Kali Linux, Ubuntu, Windows evaluation images in an isolated lab only.",
        "official_downloads": [
            ("VirtualBox", "https://www.virtualbox.org/"),
            ("Kali Linux", "https://www.kali.org/get-kali/"),
            ("Ubuntu", "https://ubuntu.com/download"),
        ],
        "refs": [
            "NIST Cybersecurity Framework — https://www.nist.gov/cyberframework",
            "OWASP — https://owasp.org/",
            "Kali documentation — https://www.kali.org/docs/",
        ],
    },
    "CLOUD_COMPUTING": {
        "title": "Beyond Cloud Computing Internship",
        "folder": "cloud_computing",
        "file_prefix": "Beyond_Cloud_Internship",
        "discipline": "cloud lab",
        "safety": "Stay within Beyond-approved free tiers. Set billing alerts. Destroy unused resources daily. Never store secrets in public repos.",
        "responsible_label": "Cloud cost & security checkpoint",
        "toolchain_note": "VirtualBox, Docker, GitHub, free tiers or local cloud-like stacks approved by Beyond.",
        "official_downloads": [
            ("Docker", "https://docs.docker.com/get-docker/"),
            ("VirtualBox", "https://www.virtualbox.org/"),
            ("Git", "https://git-scm.com/downloads"),
        ],
        "refs": [
            "Docker docs — https://docs.docker.com/",
            "NIST CSF — https://www.nist.gov/cyberframework",
        ],
    },
    "LIVE_SOUND_ENGINEERING": {
        "title": "Beyond Live Sound Engineering Internship",
        "folder": "live_sound",
        "file_prefix": "Beyond_LiveSound_Internship",
        "discipline": "live sound / AV lab",
        "safety": "Wear hearing protection during loud checks. Inspect cables. Do not overload circuits. Coil and label Beyond cables after use.",
        "responsible_label": "Hearing & electrical safety checkpoint",
        "toolchain_note": "Beyond mixers, microphones, amplifiers, speakers, and approved test tones/software.",
        "official_downloads": [
            ("Manufacturer manuals for assigned gear", "supervisor / Beyond asset docs"),
            ("Approved meter/test apps only", "supervisor list"),
        ],
        "refs": [
            "Equipment manuals assigned by Beyond",
            "Venue/electrical safety brief from supervisor",
        ],
    },
    "LIGHTING_ENGINEERING": {
        "title": "Beyond Lighting Engineering Internship",
        "folder": "lighting",
        "file_prefix": "Beyond_Lighting_Internship",
        "discipline": "lighting / stage lab",
        "safety": "Respect working-at-height rules and PPE. Calculate power before energizing. Never hang fixtures without approved hardware.",
        "responsible_label": "Rigging & power safety checkpoint",
        "toolchain_note": "Beyond fixtures/controllers and free visualizer software approved by the supervisor.",
        "official_downloads": [
            ("Approved lighting console/visualizer software", "supervisor list"),
            ("Fixture manuals", "Beyond asset library"),
        ],
        "refs": [
            "Fixture/console manuals assigned by Beyond",
            "Local electrical and height-safety rules",
        ],
    },
    "SCREENS_VIDEO": {
        "title": "Beyond Screens & Video Internship",
        "folder": "screens_video",
        "file_prefix": "Beyond_ScreensVideo_Internship",
        "discipline": "video / LED / projection lab",
        "safety": "Use approved lifts/stands. Secure cables against trip hazards. Follow power-down rules before reseating signal connectors when instructed.",
        "responsible_label": "Signal & site safety checkpoint",
        "toolchain_note": "Beyond LED screens/projectors/switchers and free media tools approved by Beyond.",
        "official_downloads": [
            ("Approved media/player tools", "supervisor list"),
            ("Equipment manuals", "Beyond asset library"),
        ],
        "refs": [
            "Equipment manuals assigned by Beyond",
            "Signal standard documentation provided in lab packs",
        ],
    },
    "INTERCOM": {
        "title": "Beyond Intercom Internship",
        "folder": "intercom",
        "file_prefix": "Beyond_Intercom_Internship",
        "discipline": "production communications lab",
        "safety": "Keep channels professional. Test backup talk paths. Never leave beltpacks unattended in public guest areas.",
        "responsible_label": "Comms etiquette & safety checkpoint",
        "toolchain_note": "Beyond wired/wireless intercom equipment and laptop configuration tools as assigned.",
        "official_downloads": [
            ("Intercom system manuals", "Beyond asset library"),
            ("Approved config utilities", "supervisor list"),
        ],
        "refs": [
            "System manuals assigned by Beyond",
            "Production call-language brief from supervisor",
        ],
    },
}

SOFTWAREISH = {
    "ARTIFICIAL_INTELLIGENCE",
    "MACHINE_LEARNING",
    "DATA_SCIENCE",
    "SOFTWARE_DEVELOPMENT",
    "CLOUD_COMPUTING",
    "NETWORKING",
    "CYBER_SECURITY",
}


def load_seed():
    return json.loads(SEED.read_text(encoding="utf-8"))


def program_tasks(seed, code: str):
    prog = next(p for p in seed["programs"] if p["code"] == code)
    rows = []
    for t in prog["tasks"]:
        title = t["title"]
        mode, topic = (title.split(":", 1) + [""])[:2] if ":" in title else ("?", title)
        rows.append(
            {
                "day": int(t["day_number"]),
                "mode": mode.strip(),
                "topic": topic.strip(),
                "difficulty": t.get("difficulty"),
                "hours": t.get("estimated_hours", 8),
                "tools": t.get("tools"),
                "objective": t.get("objective"),
                "submission": t.get("submission"),
                "instructions": t.get("instructions") or [],
                "pass_mark": t.get("pass_mark", 60),
                "study_note": t.get("study_note") or "",
            }
        )
    return prog, rows


def slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")[:48] or "Topic"


def topic_concepts(topic: str, code: str):
    """Meaningful concept rows — not word-splitting junk."""
    t = topic.lower()
    base = [
        (
            "Authorization boundary",
            "What systems and data you may touch today",
            "Complete the pre-flight checklist",
            "Working on production by mistake",
            "Stop and ask supervisor if unclear",
        ),
        (
            "Reproducible evidence",
            "Another person can rerun or inspect your work",
            "Use standard filenames and a Day folder",
            "Missing versions or secret screenshots",
            "Freeze versions; privacy-check screenshots",
        ),
        (
            "Verification",
            "A second method that confirms the result",
            "Rerun / second measurement / peer check",
            "Screenshot-only ‘proof’",
            "Pair evidence with an independent check",
        ),
    ]
    extras = []
    if any(k in t for k in ("python", "notebook", "git", "data", "machine", "neural", "classif", "regress", "cluster", "feature", "model", "prompt", "rag", "agent", "generative", "ai ", "ai orientation", "statistic", "linear")):
        extras = [
            ("Problem framing", "Who needs what, constraints, success criteria", "Write a short frame table", "Building the wrong thing", "Supervisor review of frame"),
            ("Train vs evaluation data", "Learn on one set; judge on another", "Keep a held-out TEST set", "Testing on training data", "Document the split"),
            ("Human oversight", "A person remains accountable", "Disclose AI help; verify claims", "Blind trust in outputs", "AI/tool disclosure block"),
        ]
    elif any(k in t for k in ("network", "subnet", "vlan", "routing", "osi", "tcp", "ipv", "switch", "wireless", "dns", "dhcp")):
        extras = [
            ("Layered models", "Separate physical, link, network, transport concerns", "Map today's lab to OSI/TCP-IP", "Misplacing the fault layer", "Structured troubleshooting top-down/bottom-up"),
            ("Addressing plan", "Who gets which address/mask/gateway", "Document IP plan before changes", "Duplicate IPs / wrong mask", "Written plan + ping/traceroute verify"),
            ("Change control", "Record what changed and how to roll back", "Lab journal entry", "Undocumented cable/IP edits", "Before/after notes"),
        ]
    elif any(k in t for k in ("security", "threat", "risk", "cia", "harden", "linux", "windows", "ethical", "incident", "malware", "crypto")):
        extras = [
            ("Authorization", "Written permission for the exact scope", "Confirm lab isolation", "Offensive action outside sandbox", "Written supervisor approval"),
            ("CIA triad", "Confidentiality, integrity, availability", "Map controls to CIA", "Fixing one pillar while breaking another", "Risk table in report"),
            ("Least privilege", "Only the access needed", "Use lab accounts, not admin everywhere", "Standing admin rights", "Separate admin/user sessions"),
        ]
    elif any(k in t for k in ("cloud", "docker", "container", "virtual", "kubernetes", "iam")):
        extras = [
            ("Shared responsibility", "What you secure vs what the provider secures", "List responsibilities for today's service", "Assuming the cloud is ‘safe by default’", "Checklist in report"),
            ("Cost control", "Free-tier/billing limits", "Set alerts; destroy unused resources", "Forgotten running VMs", "End-of-day teardown evidence"),
            ("Secrets handling", "Keys never in git or screenshots", "Use placeholders / env files", "Leaked tokens", "Secret scan of submission folder"),
        ]
    elif any(k in t for k in ("sound", "mic", "mixer", "gain", "audio", "speaker", "di ")):
        extras = [
            ("Signal flow", "Source → processing → amplification → speakers", "Draw today's flow", "Feedback / wrong patch", "Labelled diagram + listen test"),
            ("Gain structure", "Healthy levels without clipping/noise", "Set gain before heavy EQ", "Distorting or inaudible show", "Meter checks at each stage"),
            ("Hearing protection", "Protect hearing during loud checks", "Wear protection for loud bursts", "Permanent hearing damage", "Non-negotiable PPE rule"),
        ]
    elif any(k in t for k in ("light", "dmx", "fixture", "console", "power", "universe")):
        extras = [
            ("Power calculation", "Load vs available circuits", "Calculate before energizing", "Tripped breakers / fire risk", "Written power plan"),
            ("DMX addressing", "Unique addresses / universes", "Document address map", "Colliding fixtures", "Address sheet in evidence"),
            ("Working at height", "Approved methods and spotters", "Follow Beyond height rules", "Unsafe climb/hang", "Stop work authority"),
        ]
    elif any(k in t for k in ("video", "hdmi", "sdi", "projector", "led", "resolution", "frame")):
        extras = [
            ("Signal path", "Source → conversion → display", "Document connectors/format", "No image / wrong EDID", "Cable/format checklist"),
            ("Raster timing", "Resolution, frame rate, aspect", "Match source to display capability", "Scaling artifacts", "Confirm native timings"),
            ("Trip/lift safety", "Cable dress and approved lifts", "Secure cables; approved stands", "Injury / equipment drop", "Site safety walk"),
        ]
    elif any(k in t for k in ("intercom", "partyline", "matrix", "beltpack", "headset", "call language")):
        extras = [
            ("Channel plan", "Who talks to whom", "Write channel assignment", "Cross-talk in show-critical moment", "Posted channel card"),
            ("Call language", "Clear, short, professional phrases", "Practise standard calls", "Jokes during cues", "Etiquette checklist"),
            ("Backup path", "Second way to communicate", "Test spare beltpack/channel", "Single point of failure", "Backup test evidence"),
        ]
    else:
        extras = [
            ("Core topic skill", f"Professional practice for {topic}", "Complete today's practical objective", "Shallow buzzword-only work", "Demonstrate with evidence"),
            ("Tool discipline", "Use only approved tools/versions", "Verify tools before lab", "Unapproved downloads", "Official sources only"),
            ("Professional documentation", "Clear enough for handover", "Write procedure + results", "Undocumented changes", "README + report template"),
        ]
    return extras + base


def mode_schedule(mode: str, hours: float):
    h = float(hours or 8)
    if mode == "Orientation and setup":
        return [
            ("0:00–0:25", "Review prior day + authorization checklist", "Signed pre-flight"),
            ("0:25–1:10", "Concept study for today's topic", "Concept notes"),
            ("1:10–2:40", "Environment verify / light setup lab", "Working baseline"),
            ("2:40–2:55", "Break", "—"),
            ("2:55–4:40", "Guided mini-exercise + expected checks", "First evidence"),
            ("4:40–4:55", "Break", "—"),
            ("4:55–6:25", "Verification + troubleshooting notes", "Rerun proof"),
            ("6:25–8:00", "Evidence pack + report draft", "Submission draft"),
        ]
    if mode == "Guided practical":
        return [
            ("0:00–0:20", "Authorization + tool verify", "Pre-flight"),
            ("0:20–1:50", "Guided practical steps 1–N", "Working artifact"),
            ("1:50–2:05", "Break", "—"),
            ("2:05–3:50", "Evaluation / measurement", "Metrics or checks"),
            ("3:50–4:05", "Break", "—"),
            ("4:05–5:50", "Responsible checkpoint + fixes", "Risk table"),
            ("5:50–7:10", "Evidence + report", "Pack"),
            ("7:10–8:00", "Self-check + supervisor rehearsal", "Ready pack"),
        ]
    if mode == "Independent build":
        return [
            ("0:00–0:30", "Rewrite objective in own words", "Frame table"),
            ("0:30–2:30", "Independent build", "Artifact"),
            ("2:30–2:45", "Break", "—"),
            ("2:45–4:15", "Test / measure / review", "Results"),
            ("4:15–5:45", "Documentation while building", "README notes"),
            ("5:45–6:00", "Break", "—"),
            ("6:00–8:00", "Evidence + report + rehearsal", "Ready pack"),
        ]
    if mode == "Troubleshoot":
        return [
            ("0:00–0:25", "Confirm healthy baseline", "Baseline evidence"),
            ("0:25–2:00", "Fault A diagnose → fix", "Fault log A"),
            ("2:00–2:15", "Break", "—"),
            ("2:15–3:45", "Fault B diagnose → fix", "Fault log B"),
            ("3:45–5:00", "Independent verification", "Rerun proof"),
            ("5:00–6:30", "Incident-style report", "Report"),
            ("6:30–8:00", "Five-minute failure drill + submit", "Ready pack"),
        ]
    if mode == "Document and improve":
        return [
            ("0:00–0:40", "Audit prior artifacts for gaps", "Gap list"),
            ("0:40–2:20", "Rewrite procedure for reproducibility", "README"),
            ("2:20–2:35", "Break", "—"),
            ("2:35–4:35", "One improvement + before/after proof", "Improvement evidence"),
            ("4:35–6:20", "Report + packaging", "Pack"),
            ("6:20–8:00", "Peer-style checklist + rehearsal", "Ready pack"),
        ]
    # Assessment
    return [
        ("0:00–0:30", "Assessment brief + rubric review", "Plan"),
        ("0:30–2:30", "End-to-end demonstration", "Demo evidence"),
        ("2:30–2:45", "Break", "—"),
        ("2:45–4:15", "Self-score against rubric", "Self-score sheet"),
        ("4:15–5:45", "Reflection: wins, failures, next steps", "Reflection"),
        ("5:45–8:00", "Final evidence pack + report", "Submission"),
    ]


def lab_code(day: dict, code: str) -> str:
    d = day["day"]
    topic = day["topic"]
    mode = day["mode"]
    if code in SOFTWAREISH:
        return f'''# day_{d:02d}_lab.py — {mode}: {topic}
# Beyond internship lab starter. Synthetic/approved data only. Not production-ready.

"""Day {d} practical starter for {topic}.

Instructions for the student:
1. Keep this file under src/.
2. Replace TODO sections with your work for today's mode.
3. Print clear expected outputs.
4. Capture evidence filenames listed in the handbook.
"""

from __future__ import annotations

def main() -> None:
    print("Beyond internship lab")
    print("Program topic:", {topic!r})
    print("Day mode:", {mode!r})
    # TODO: implement today's practical checks below.
    checks = {{
        "authorization_confirmed": True,  # set False until checklist complete
        "environment_verified": True,
        "practical_step_done": False,  # flip to True when your lab step works
    }}
    for k, v in checks.items():
        print(f"{{k}}: {{v}}")
    if not all(checks.values()):
        print("INCOMPLETE: finish TODO practical work, then rerun.")
        return
    print("OK: starter checks passed. Replace TODOs with real measurements/outputs.")

if __name__ == "__main__":
    main()
'''
    return f"""# Day {d} lab checklist script (text companion)
# Topic: {topic}
# Mode: {mode}
# This program is equipment-led. Use the handbook steps on Beyond-authorized gear.
# Mark each line DONE in your report when complete. Do not invent equipment access.

[ ] Authorization and PPE confirmed
[ ] Equipment / software from the approved tools list verified
[ ] Practical procedure for "{topic}" completed
[ ] Independent verification (second check / peer / supervisor) completed
[ ] Evidence captured with day-{d:02d}_ filenames
[ ] Safety controls recorded
"""


def ensure_styles(doc: Document):
    for name, color in {
        "Beyond Note": BLUE,
        "Beyond Safety": WARN,
        "Beyond Warning": WARN,
    }.items():
        try:
            doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, 1)
            st.font.size = Pt(10)
            st.font.color.rgb = color
            st.font.bold = True


def set_run(run, *, bold=False, size=11, color=None, name=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if name:
        run.font.name = name


def add_p(doc, text="", *, bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run(r, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_callout(doc, kind: str, title: str, body: str):
    style = {"note": "Beyond Note", "safety": "Beyond Safety", "warning": "Beyond Warning"}.get(kind, "Beyond Note")
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    r = p.add_run(f"{title}  {body}")
    set_run(r, bold=True, size=10, color=BLUE if kind == "note" else WARN)
    p.paragraph_format.space_after = Pt(8)


def add_code(doc, text: str):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line if line else " ")
        set_run(r, size=9, name="Consolas")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F1F5F9")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            set_run(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for r in p.runs:
            set_run(r, size=11)


def shade_header(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0B3F90")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(9)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    shade_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def new_doc(code: str) -> Document:
    if code == "ARTIFICIAL_INTELLIGENCE" and AI_DAY001.exists():
        doc = Document(str(AI_DAY001))
        clear_body(doc)
    else:
        doc = Document()
    ensure_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    return doc


def one_liner(day: dict) -> str:
    return (
        f"Complete the {day['mode'].lower()} for {day['topic']}: "
        f"{(day.get('objective') or 'meet the curriculum objective safely and document evidence.').rstrip('.')}."
    )


def build_handbook(code: str, day: dict, out_dir: Path, *, skip_existing_approved_ai_day1=True) -> Path | None:
    meta = PROGRAM_META[code]
    dnum = day["day"]
    fname = f"{meta['file_prefix']}_Day_{dnum:03d}_Student_Handbook.docx"
    out = out_dir / fname

    # Never overwrite canonical AI Day 001 reference
    if code == "ARTIFICIAL_INTELLIGENCE" and dnum == 1 and skip_existing_approved_ai_day1:
        ref = out_dir / fname
        if ref.exists() or AI_DAY001.exists():
            return None

    doc = new_doc(code)
    topic = day["topic"]
    mode = day["mode"]
    folder = f"Day_{dnum:02d}_{slug(topic)}"
    concepts = topic_concepts(topic, code)
    schedule = mode_schedule(mode, day.get("hours") or 8)
    software = code in SOFTWAREISH

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["title"].upper())
    set_run(r, bold=True, size=14, color=BLUE)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"DAY {dnum} STUDENT HANDBOOK")
    set_run(r, bold=True, size=22, color=BLUE)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(topic)
    set_run(r, bold=True, size=14, color=DARK)

    add_p(doc, f"{mode}. Full student study guide and practical laboratory manual for Beyond Enterprise interns.")
    add_table(
        doc,
        ["DAY MODE", "LEVEL", "DURATION", "CORE TOOLS"],
        [[mode, day.get("difficulty") or "foundation", f"{day.get('hours') or 8} hours", (day.get("tools") or "")[:90]]],
    )
    add_callout(doc, "note", f"Day {dnum} objective", one_liner(day))
    add_callout(doc, "safety", "Non-negotiable rule", meta["safety"])
    if dnum > 1:
        add_callout(
            doc,
            "warning",
            "Continuity",
            "Day 1 safety, evidence, and documentation rules remain in force. "
            + meta["toolchain_note"],
        )

    doc.add_heading("How to use this handbook", level=1)
    add_p(
        doc,
        "Follow the activities in order. Capture evidence at checkpoints. Write the report while the work is fresh. "
        "Do not skip verification. Stop and ask a supervisor before changing production systems, shared accounts, "
        "or using unapproved cloud/AI services.",
    )
    add_bullets(
        doc,
        [
            "Order: authorization → tool verify → lab → independent verification → evidence → report → checklists.",
            f"Submit: {(day.get('submission') or 'evidence pack + technical report')}.",
            "Supervisor approval required for production changes, paid services, or offensive/security-sensitive actions outside the sandbox.",
        ],
    )

    doc.add_heading("Suggested daily schedule", level=1)
    add_table(doc, ["TIME", "ACTIVITY", "OUTPUT"], schedule)

    doc.add_heading("Learning outcomes", level=1)
    add_p(doc, f"Curriculum objective (source of truth): {day.get('objective')}", size=10)
    add_bullets(
        doc,
        [
            f"Explain {topic} to a supervisor in under two minutes.",
            f"Configure or verify the approved tools required for {topic}.",
            f"Implement or perform the {mode.lower()} practical for this topic.",
            "Verify the result with at least one independent method.",
            "Diagnose a common failure using the troubleshooting table.",
            "Document procedure, results, limitations, and safety controls professionally.",
        ],
    )

    doc.add_heading("Key concepts", level=1)
    add_p(
        doc,
        f"Define the professional meaning of “{topic}” before you use tools. "
        "Do not split the title into meaningless words. Connect each concept to today's practical action.",
    )
    add_table(
        doc,
        ["TERM", "PLAIN-LANGUAGE MEANING", "TODAY'S ACTION", "COMMON FAILURE", "CONTROL"],
        [list(row) for row in concepts],
    )

    doc.add_heading("Prerequisites and authorization check", level=1)
    add_bullets(
        doc,
        [
            f"Prior learning: complete earlier days in this program through Day {max(1, dnum-1)} as assigned.",
            "Authorized Beyond workstation, lab, VM, or equipment only.",
            "Required permissions confirmed (admin install rights only if approved).",
            f"Tools available: {(day.get('tools') or 'as assigned')}.",
            "Backup/rollback plan: do not overwrite prior Day evidence folders.",
            "Data/equipment classification understood; no unauthorized sensitive data.",
        ],
    )
    add_callout(doc, "safety", "Authorization", "If any item is unclear, stop and ask the supervisor. Do not bypass security controls.")

    doc.add_heading("Download and installation guide", level=1)
    if dnum == 1:
        add_p(doc, "Day 1 introduces or verifies the program toolchain. Use official sources only.")
        add_table(
            doc,
            ["TOOL", "OFFICIAL SOURCE"],
            [[n, u] for n, u in meta["official_downloads"]],
        )
        if software:
            add_p(doc, "Platform notes:", bold=True)
            add_bullets(
                doc,
                [
                    "Windows 10/11: use official installers; prefer py -m pip after Python install.",
                    "macOS (Apple silicon/Intel): use official universal/supported builds; do not delete system Python.",
                    "Ubuntu/Debian: use apt for system packages; prefer venv/containers for project tools; ask if sudo is blocked.",
                ],
            )
            add_code(
                doc,
                """# Version verification examples
python3 --version || py --version
git --version
code --version""",
            )
        else:
            add_p(
                doc,
                "This program is equipment-led. Confirm serials/asset tags, power, cabling, and PPE with the supervisor. "
                "Install only supervisor-approved companion software from official vendor channels.",
            )
    else:
        add_p(
            doc,
            "Do not repeat a full Day-1 install unless verification fails and the supervisor approves. "
            "Verify the tools listed for this day, then continue.",
        )
        add_table(doc, ["TOOL / CHECK", "VERIFICATION"], [[(day.get("tools") or "Assigned tools"), "Confirm presence/version with supervisor list"]])
        if software:
            add_code(
                doc,
                """python3 --version || py --version
python3 -m pip check || py -m pip check""",
            )

    doc.add_heading("Project and file organization", level=1)
    prog_root = meta["file_prefix"]
    if software:
        add_code(
            doc,
            f"""{prog_root}/
└── {folder}/
    ├── evidence/
    ├── notebooks/
    ├── src/
    ├── data/
    ├── output/
    ├── report/
    ├── requirements.txt
    └── README.md""",
        )
        add_p(doc, "Windows PowerShell:", bold=True)
        add_code(
            doc,
            f'cd "$HOME\\Documents"\nmkdir {prog_root}\\{folder}\\evidence, {prog_root}\\{folder}\\notebooks, {prog_root}\\{folder}\\src, {prog_root}\\{folder}\\data, {prog_root}\\{folder}\\output, {prog_root}\\{folder}\\report -Force',
        )
        add_p(doc, "macOS / Linux:", bold=True)
        add_code(
            doc,
            f"cd ~/Documents\nmkdir -p {prog_root}/{folder}/{{evidence,notebooks,src,data,output,report}}",
        )
    else:
        add_code(
            doc,
            f"""{prog_root}/
└── {folder}/
    ├── evidence/
    ├── diagrams/
    ├── checklists/
    ├── report/
    └── README.md""",
        )
    add_callout(doc, "warning", "Secrets", "Never place passwords, API keys, tokens, personal records, or client data in the submission folder.")

    doc.add_heading("Environment setup", level=1)
    if software:
        add_numbered(
            doc,
            [
                "Open today's folder in VS Code (or approved editor).",
                "Activate the project virtual environment if this program uses Python.",
                "Select the correct interpreter/kernel.",
                "Install only packages required for this day (supervisor-approved).",
                "Export requirements when asked: python -m pip freeze > requirements.txt",
                "Use placeholders such as API_KEY=YOUR_APPROVED_LAB_KEY — never real secrets.",
            ],
        )
    else:
        add_numbered(
            doc,
            [
                "Inspect assigned Beyond equipment and power conditions.",
                "Confirm PPE and spotter requirements.",
                "Stage cables/connectors/tools on a clean, labelled work area.",
                "Photograph the starting state for evidence (no private backgrounds).",
                "Confirm rollback/restore condition with the supervisor before major changes.",
            ],
        )

    doc.add_heading("Step-by-step practical lab", level=1)
    add_p(doc, f"{mode}: {topic}", bold=True, size=12, color=BLUE)
    add_p(doc, "Curriculum instructions (expanded below into executable steps):")
    for line in day.get("instructions") or []:
        add_p(doc, f"• {line}", size=10)

    if software:
        add_p(doc, f"Create `src/day_{dnum:02d}_lab.py` (starter):")
        add_code(doc, lab_code(day, code))
    else:
        add_p(doc, "Equipment checklist companion (copy into report and mark DONE):")
        add_code(doc, lab_code(day, code))

    steps = [
        ("Pre-flight authorization", "Complete checklist", "All boxes checked", "Safe scope", f"day-{dnum:02d}_preflight.png", "Unclear permission", "Stop; ask supervisor"),
        ("Prepare workspace / project folder", "Create Day folder tree", "Folder exists", "Organization", f"day-{dnum:02d}_folders.png", "Wrong location", "Use Documents path / approved lab path"),
        ("Verify tools", "Version or equipment check", "Tools confirmed", "Readiness", f"day-{dnum:02d}_tools.png", "Missing tool", "Use official source or ask supervisor"),
        ("Execute practical", "Follow mode procedure for topic", "Expected lab result", "Skill demonstration", f"day-{dnum:02d}_practical.png", "Procedure fails", "Change one variable; log hypothesis"),
        ("Independent verification", "Second method / rerun / peer", "Confirmed result", "Proof beyond a screenshot", f"day-{dnum:02d}_verify.png", "Mismatch", "Do not fabricate; document and escalate if stuck"),
        ("Package evidence + report", "Filenames + report template", "Complete pack", "Assessable submission", f"day-{dnum:02d}_pack_list.png", "Missing files", "Use submission checklist"),
    ]
    for i, (action, cmd, exp, proves, evid, fail, fix) in enumerate(steps, 1):
        doc.add_heading(f"Step {i}: {action}", level=2)
        add_table(
            doc,
            ["FIELD", "DETAIL"],
            [
                ["Action", action],
                ["Command / work", cmd],
                ["Expected result", exp],
                ["What it proves", proves],
                ["Evidence", evid],
                ["Likely failure", fail],
                ["Safe corrective action", fix],
            ],
        )

    doc.add_heading(meta["responsible_label"], level=1)
    add_table(
        doc,
        ["RISK", "POSSIBLE HARM", "PREVENTION", "DETECTION", "APPROVER"],
        [
            ["Unauthorized system/data use", "Breach, outage, legal issue", "Lab-only scope", "Supervisor spot check", "Supervisor"],
            ["Unsafe procedure", "Injury or equipment damage", "PPE + written steps", "Buddy/supervisor observe", "Student + supervisor"],
            ["Secret/PII leakage", "Privacy incident", "Synthetic/approved data; redact screenshots", "Submission review", "Student"],
            ["Irreproducible work", "Cannot assess internship", "Versions + rerun proof", "Second-run evidence", "Student"],
            ["Overclaiming readiness", "Bad production decision", "State limitations clearly", "Rubric / oral check", "Supervisor"],
        ],
    )

    doc.add_heading("Evaluation and metrics", level=1)
    if software and any(k in topic.lower() for k in ("classif", "model", "regress", "cluster", "evaluat", "ml", "neural", "recommend")):
        add_table(
            doc,
            ["METRIC", "QUESTION", "CAUTION"],
            [
                ["Accuracy / error rate", "Overall correctness?", "Weak on imbalanced data"],
                ["Precision / recall / F1", "Error type trade-off?", "Optimize for the real cost of mistakes"],
                ["Latency / resource use", "Is it usable?", "Lab laptop ≠ production load"],
            ],
        )
    else:
        add_table(
            doc,
            ["CHECK", "SUCCESS CRITERION", "LIMITATION"],
            [
                ["Functional result", "Meets today's objective", "Classroom demo ≠ production"],
                ["Safety compliance", "No rule breached", "Still needs supervisor judgment"],
                ["Reproducibility", "Second check agrees", "Environment drift possible"],
            ],
        )

    doc.add_heading("Independent verification", level=1)
    add_bullets(
        doc,
        [
            "Rerun the procedure after restarting the tool/session/environment.",
            "Use a second method: manual calculation, alternate meter/path, unit test, log review, or peer check.",
            "Distinguish evidence (what you captured) from proof (why it is correct).",
        ],
    )

    doc.add_heading("Troubleshooting guide", level=1)
    add_table(
        doc,
        ["SYMPTOM", "LIKELY CAUSE", "SAFE FIRST ACTION", "VERIFY AFTER", "ESCALATE IF"],
        [
            ["Tool missing / won't start", "Wrong install or PATH", "Official reinstall/verify only if approved", "Version command works", "Image is locked by IT"],
            ["Lab step fails", "Skipped prerequisite", "Re-read prior step; change one variable", "Expected output appears", "Hardware fault suspected"],
            ["Intermittent result", "Cable/power/network flake", "Reseat/inspect; simplify path", "Stable for 3 repeats", "Shared infra involved"],
            ["Permission error", "Account/policy limit", "Do not bypass; ask supervisor", "Approved alternate path", "Blocks required evidence"],
            ["Unclear authorization", "Scope ambiguity", "STOP work", "Written clarification", "Always escalate"],
        ],
    )
    add_callout(
        doc,
        "warning",
        "Five-minute failure procedure",
        "Before review: correct folder, tools verified, rerun practical, open evidence files, hide notifications/secrets, know your main limitation.",
    )

    doc.add_heading("Evidence checkpoints", level=1)
    add_table(
        doc,
        ["FILE", "MUST SHOW", "WHY", "PRIVACY/SAFETY", "FORMAT"],
        [
            [f"day-{dnum:02d}_tools.png", "Tools/versions or staged gear", "Readiness", "No secrets", "PNG/JPEG"],
            [f"day-{dnum:02d}_practical.png", "Key practical result", "Execution", "No PII", "PNG/JPEG"],
            [f"day-{dnum:02d}_verify.png", "Second check", "Verification", "No notifications", "PNG/JPEG"],
            [f"report/day-{dnum:02d}_technical_report.md", "Full report", "Assessment", "No credentials", "MD/PDF"],
            ["README.md", "How to reproduce", "Handover", "No secrets", "MD"],
        ],
    )

    doc.add_heading("Professional technical report guide", level=1)
    add_numbered(
        doc,
        [
            "Title and metadata",
            "Objective",
            "Scope and authorization",
            "Concept summary",
            "Environment / equipment and versions",
            "Data or input description",
            "Procedure",
            "Results",
            "Evaluation",
            "Evidence references",
            "Safety, ethics, and risk controls",
            "Independent verification",
            "Troubleshooting",
            "Limitations",
            "Conclusion and next step",
            "AI/tool assistance disclosure (if any)",
            "References",
        ],
    )
    add_callout(
        doc,
        "note",
        "Example AI disclosure",
        "I used [approved tool/model] to assist with [specific task]. I did not provide personal or client data. "
        "I independently verified [claims/results] using [method]. I corrected or rejected [brief description], "
        "and I remain responsible for the submitted work.",
    )

    doc.add_heading("Submission checklists", level=1)
    doc.add_heading("A. Safety and ethics checklist", level=2)
    add_bullets(doc, ["Authorized scope only", "PPE/lab rules followed", "No secrets/PII in submission", "Honest time and troubleshooting notes", meta["safety"][:120] + "…"])
    doc.add_heading("B. Technical completion checklist", level=2)
    add_bullets(doc, [f"Objective addressed for {topic}", f"{mode} practical completed", "Verification done", "Tools/equipment checked", "Limitations stated"])
    doc.add_heading("C. Evidence and report checklist", level=2)
    add_bullets(doc, ["Required screenshots/files present", "Filenames follow day-XX_… pattern", "Report complete", f"Matches: {(day.get('submission') or 'evidence + report')}", "Prior day evidence not overwritten"])

    doc.add_heading("Self-check questions", level=1)
    add_numbered(
        doc,
        [
            f"Explain {topic} without reading.",
            "What method did you choose and why?",
            "How did you verify the result?",
            "Name one failure mode and its control.",
            "What is the first safe troubleshooting action if the demo fails?",
            "Why is today's work not production-ready?",
            "What requires supervisor approval before you go further?",
        ],
    )

    doc.add_heading("Two-minute supervisor explanation", level=1)
    add_p(
        doc,
        f"Today is Day {dnum} ({mode}) on {topic}. I worked only in the authorized Beyond {meta['discipline']}. "
        f"I completed the practical objective, verified with a second check, and packaged evidence. "
        f"Main risk: unauthorized or unsafe action; control: checklist + supervisor escalation. "
        f"Limitation: classroom/lab scope only. Next: continue the topic cycle and keep documentation reproducible.",
    )

    doc.add_heading("Supervisor assessment rubric", level=1)
    add_table(
        doc,
        ["AREA", "PASS EVIDENCE", "WEIGHT"],
        [
            ["Safe and authorized setup", "Checklist + lab-only scope", "15%"],
            ["Technical execution", "Working practical result", "25%"],
            ["Concept understanding", "Clear oral/written definitions", "15%"],
            ["Evaluation and verification", "Second-method proof", "15%"],
            ["Responsible practice", "Risk controls applied", "10%"],
            ["Evidence quality", "Clear files, no secrets", "10%"],
            ["Professional reporting", "Complete report", "5%"],
            ["Honest troubleshooting", "Real notes / no fabrication", "5%"],
        ],
    )
    add_p(doc, f"Pass mark from curriculum: {day.get('pass_mark', 60)}%.")

    doc.add_heading("Official references", level=1)
    add_p(doc, f"Access check date: {date.today().isoformat()}. Prefer first-party sources. Do not invent URLs.")
    add_bullets(doc, meta["refs"])
    add_callout(
        doc,
        "note",
        "Instructor Review Required",
        "If local Beyond policy conflicts with an optional tool or online service mentioned here, follow Beyond policy and record the Assumption in your report.",
    )

    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def parse_days(spec: str):
    if "-" in spec and "," not in spec:
        a, b = spec.split("-", 1)
        return list(range(int(a), int(b) + 1))
    return [int(x) for x in spec.split(",") if x.strip()]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--program", help="Program code, e.g. ARTIFICIAL_INTELLIGENCE")
    ap.add_argument("--all-programs", action="store_true")
    ap.add_argument("--days", default="1-180")
    ap.add_argument("--force-day1", action="store_true", help="Allow regenerating AI day 1 (default: skip)")
    args = ap.parse_args()

    seed = load_seed()
    codes = [p["code"] for p in seed["programs"]] if args.all_programs else [args.program]
    if not codes or codes == [None]:
        raise SystemExit("Provide --program CODE or --all-programs")

    day_list = parse_days(args.days)
    total = 0
    skipped = 0
    for code in codes:
        if code not in PROGRAM_META:
            print("Unknown program", code)
            continue
        meta = PROGRAM_META[code]
        out_dir = HERE / meta["folder"]
        _, days = program_tasks(seed, code)
        by_day = {d["day"]: d for d in days}
        # export compact for convenience
        (out_dir).mkdir(parents=True, exist_ok=True)
        (out_dir / f"_{code.lower()}_days_compact.json").write_text(json.dumps(days, indent=2), encoding="utf-8")
        for n in day_list:
            if n not in by_day:
                continue
            path = build_handbook(
                code,
                by_day[n],
                out_dir,
                skip_existing_approved_ai_day1=not args.force_day1,
            )
            if path is None:
                skipped += 1
                print(f"SKIP {code} day {n} (canonical day 1 preserved)")
                continue
            total += 1
            if total % 20 == 0:
                print(f"... {total} handbooks written (latest {path.name})")
        # changelog
        cl = out_dir / "CHANGELOG.md"
        with cl.open("a", encoding="utf-8") as f:
            f.write(f"\n## {date.today().isoformat()} — generated days {day_list[0]}-{day_list[-1]} for {code}\n")
            f.write(f"- Wrote/updated handbooks in `{out_dir.name}/`\n")
        print(f"Done {code}: batch complete")
    print(f"TOTAL written: {total}, skipped: {skipped}")


if __name__ == "__main__":
    main()
