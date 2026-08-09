#!/usr/bin/env python3
"""
Generate Beyond AI Internship Student Handbook DOCX files (Days 2+).
Uses Day 001 as style/template reference; does not overwrite Day 001.
"""

from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SEED = HERE.parents[1] / "beyond_180_day_curriculum_seed.json"
DAY001 = HERE / "Beyond_AI_Internship_Day_001_Student_Handbook.docx"
COMPACT = HERE / "_ai_days_compact.json"
INDEX_DOCX = HERE / "Beyond_AI_Internship_Handbook_Master_Index.docx"
CHANGELOG = HERE / "CHANGELOG.md"

BLUE = RGBColor(0x0B, 0x3F, 0x90)
DARK = RGBColor(0x1E, 0x29, 0x3B)


def load_day(day_num: int) -> dict:
    days = json.loads(COMPACT.read_text())
    return next(d for d in days if d["day"] == day_num)


def slug_topic(topic: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", topic).strip("_")
    return s[:48] or "Topic"


def ensure_styles(doc: Document):
    """Ensure Beyond custom paragraph styles exist (clone from Day 001 if present)."""
    needed = {
        "Beyond Note": RGBColor(0x0B, 0x3F, 0x90),
        "Beyond Safety": RGBColor(0x9A, 0x34, 0x12),
        "Beyond Warning": RGBColor(0x9A, 0x34, 0x12),
    }
    styles = doc.styles
    for name, color in needed.items():
        try:
            styles[name]
        except KeyError:
            style = styles.add_style(name, 1)  # paragraph
            style.font.size = Pt(10)
            style.font.color.rgb = color
            style.font.bold = True


def set_run(run, *, bold=False, size=11, color=None, name=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if name:
        run.font.name = name


def add_p(doc, text="", *, style="Normal", bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph(style=style if style in [s.name for s in doc.styles] else "Normal")
    # If custom style missing, fall back and format run
    if not text:
        return p
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_callout(doc, kind: str, title: str, body: str):
    style_map = {
        "note": "Beyond Note",
        "safety": "Beyond Safety",
        "warning": "Beyond Warning",
    }
    style = style_map.get(kind, "Beyond Note")
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    run = p.add_run(f"{title}  {body}")
    set_run(run, bold=True, size=10, color=BLUE if kind == "note" else RGBColor(0x9A, 0x34, 0x12))
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code(doc, text: str):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        set_run(run, size=9, name="Consolas")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # light shade
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F1F5F9")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            set_run(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for r in p.runs:
            set_run(r, size=11)


def shade_header(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0B3F90")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(9)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    shade_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def new_doc_from_template() -> Document:
    if DAY001.exists():
        doc = Document(str(DAY001))
        clear_body(doc)
    else:
        doc = Document()
    ensure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    return doc


def day_folder(day: dict) -> str:
    return f"Day_{day['day']:02d}_{slug_topic(day['topic'])}"


def mode_lab_spec(day: dict) -> dict:
    """Return mode-specific lab narrative and commands for Days 2–5 (topic: AI Orientation)."""
    mode = day["mode"]
    d = day["day"]
    folder = day_folder(day)
    notebook = f"day_{d:02d}_responsible_ai_lab.ipynb"
    script = f"day_{d:02d}_lab.py"

    common_verify = [
        "python --version",
        "python -c \"import sys; print(sys.executable)\"",
        "python -m pip check",
    ]

    if mode == "Guided practical":
        return {
            "one_liner": "Complete a guided responsible-AI laboratory that builds on Day 1: problem frame, heuristic baseline, evaluation, bias/privacy check, and verified evidence pack.",
            "safety": "Never paste client, patient, student, employee, financial, credential, or other private data into a public AI service. Use only synthetic Day-2 lab data.",
            "schedule": [
                ("0:00–0:20", "Review Day 1 outcomes + authorization checklist", "Signed pre-flight checklist"),
                ("0:20–0:40", "Verify Python/venv/VS Code/Jupyter (no reinstall)", "Version screenshot"),
                ("0:40–1:40", "Guided notebook: problem frame + heuristic", "Working notebook cells"),
                ("1:40–1:55", "Break", "—"),
                ("1:55–3:10", "Evaluation metrics + bias/privacy checkpoint", "Metric table + risk table"),
                ("3:10–4:10", "Prompting exercise (synthetic only) + independent verification", "Second-run proof"),
                ("4:10–4:25", "Break", "—"),
                ("4:25–6:10", "Evidence packaging + technical report draft", "Evidence folder + report draft"),
                ("6:10–7:10", "Troubleshooting practice + finalize submission", "Troubleshoot notes"),
                ("7:10–8:00", "Self-check + two-minute supervisor rehearsal", "Ready submission pack"),
            ],
            "outcomes": [
                "Explain AI orientation and responsible use without reading notes.",
                "Configure and verify an existing Day-1 Python virtual environment.",
                "Implement a small deterministic heuristic on synthetic data.",
                "Measure accuracy/precision/recall on a held-out synthetic set.",
                "Diagnose at least one common lab failure using the troubleshooting table.",
                "Document AI assistance, verification, and limitations in a professional report.",
            ],
            "lab_title": "Guided responsible-AI laboratory",
            "notebook": notebook,
            "script": script,
            "folder": folder,
            "verify_cmds": common_verify,
            "install_needed": False,
            "code": '''# day_02_lab.py — Guided practical: responsible AI baseline
# Synthetic data only. Deterministic. Not production-ready.

from dataclasses import dataclass
from typing import List, Tuple

@dataclass(frozen=True)
class Ticket:
    text: str
    urgent_truth: bool  # ground truth for evaluation only

# Synthetic helpdesk-like tickets (NOT real customer data)
TRAIN = [
    Ticket("server down cannot login", True),
    Ticket("please reset my password", False),
    Ticket("payment failed urgent", True),
    Ticket("how do I export a report", False),
    Ticket("database unreachable production", True),
    Ticket("change my display name", False),
]

TEST = [
    Ticket("prod api timeout critical", True),
    Ticket("where is the user guide", False),
    Ticket("cannot access payroll system", True),
    Ticket("update my email signature", False),
]

URGENT_KEYWORDS = {"down", "urgent", "critical", "unreachable", "timeout", "payroll", "payment"}

def predict_urgent(text: str) -> bool:
    """Heuristic: keyword match. Transparent and editable by humans."""
    tokens = set(text.lower().split())
    return len(tokens & URGENT_KEYWORDS) > 0

def evaluate(rows: List[Ticket]) -> Tuple[float, float, float, float]:
    tp = fp = tn = fn = 0
    for r in rows:
        pred = predict_urgent(r.text)
        if pred and r.urgent_truth: tp += 1
        elif pred and not r.urgent_truth: fp += 1
        elif (not pred) and r.urgent_truth: fn += 1
        else: tn += 1
    accuracy = (tp + tn) / max(tp + tn + fp + fn, 1)
    precision = tp / max(tp + fp, 1)
    recall = tp / max(tp + fn, 1)
    f1 = 0.0 if precision + recall == 0 else 2 * precision * recall / (precision + recall)
    return accuracy, precision, recall, f1

if __name__ == "__main__":
    # Smoke test on TRAIN (for learning only — report TEST metrics as primary)
    acc_tr, p_tr, r_tr, f1_tr = evaluate(TRAIN)
    acc, p, r, f1 = evaluate(TEST)
    print("TRAIN (not primary):", round(acc_tr, 3), round(p_tr, 3), round(r_tr, 3), round(f1_tr, 3))
    print("TEST  (primary):    ", round(acc, 3), round(p, 3), round(r, 3), round(f1, 3))
    assert acc >= 0.5, "Unexpectedly weak heuristic — check keyword set / labels"
    print("OK: heuristic ran and was independently evaluated on TEST.")
''',
            "expected": "TEST metrics print; script ends with OK; no secrets in output.",
            "fault_for_day4": None,
        }

    if mode == "Independent build":
        return {
            "one_liner": "Independently design a small responsible-AI decision aid on synthetic data, without copying Day 2 line-for-line, and defend your design choices.",
            "safety": "Do not scrape real tickets or use personal messages. Keep the system offline or on localhost only.",
            "schedule": [
                ("0:00–0:25", "Re-state problem frame in your own words", "Problem-frame table"),
                ("0:25–0:45", "Environment verification", "Version evidence"),
                ("0:45–2:15", "Independent build of heuristic or tiny classifier", "Working src + notebook"),
                ("2:15–2:30", "Break", "—"),
                ("2:30–4:00", "Evaluation on held-out synthetic set + error analysis", "Metric + confusion notes"),
                ("4:00–5:00", "Responsible-AI checkpoint + controls", "Risk table"),
                ("5:00–5:15", "Break", "—"),
                ("5:15–6:45", "Evidence + report", "Submission pack"),
                ("6:45–8:00", "Self-check + supervisor rehearsal", "Ready pack"),
            ],
            "outcomes": [
                "Frame an AI-relevant problem with scope, constraints, and success criteria.",
                "Implement an independent solution (heuristic preferred unless you can justify ML).",
                "Compare at least two design choices and defend the selected method.",
                "Verify results with a second method (rerun + manual spot-check).",
                "Document limitations and why the work is not production-ready.",
            ],
            "lab_title": "Independent responsible-AI mini-build",
            "notebook": notebook,
            "script": script,
            "folder": folder,
            "verify_cmds": common_verify,
            "install_needed": False,
            "code": '''# day_03_lab.py — Independent build scaffold (replace thresholds/keywords with YOUR design)
# Synthetic only. You must change KEYWORDS / rules and justify them in the report.

from dataclasses import dataclass

@dataclass(frozen=True)
class Case:
    text: str
    label: bool

# TODO: replace with your own synthetic cases (minimum 12 total; keep a held-out TEST)
CASES = [
    Case("billing error on invoice", True),
    Case("thanks for the update", False),
    Case("security alert login from new device", True),
    Case("add me to the mailing list", False),
    Case("shipment delayed critical", True),
    Case("office wifi password reminder", False),
]

KEYWORDS = {"billing", "security", "critical", "alert", "delayed"}  # CHANGE these

def predict(text: str) -> bool:
    return any(k in text.lower() for k in KEYWORDS)

def accuracy(rows):
    correct = sum(predict(r.text) == r.label for r in rows)
    return correct / max(len(rows), 1)

if __name__ == "__main__":
    # Simple split: first 4 train-like, last 2 test-like (improve this yourself)
    test = CASES[-2:]
    print("TEST accuracy:", round(accuracy(test), 3))
    for r in test:
        print(" ", r.text, "->", predict(r.text), "(truth", r.label, ")")
    print("OK: independent scaffold executed. Improve data split and rules before submit.")
''',
            "expected": "Script runs; you customized keywords/data; TEST accuracy reported; design defended in report.",
            "fault_for_day4": None,
        }

    if mode == "Troubleshoot":
        return {
            "one_liner": "Diagnose and repair controlled failures in the responsible-AI lab environment and heuristic, recording an honest hypothesis → test → result loop.",
            "safety": "Do not disable security controls, delete failure evidence, or invent successful metrics. Escalate if authorization is unclear.",
            "schedule": [
                ("0:00–0:20", "Baseline: confirm Day 2/3 artifacts still run", "Baseline screenshot"),
                ("0:20–1:20", "Inject/accept Fault A (env) and diagnose", "Fault log A"),
                ("1:20–1:35", "Break", "—"),
                ("1:35–3:05", "Fault B (logic/metrics) diagnose + fix", "Fault log B"),
                ("3:05–4:05", "Independent verification after repairs", "Rerun proof"),
                ("4:05–4:20", "Break", "—"),
                ("4:20–6:20", "Evidence + incident-style report", "Report"),
                ("6:20–8:00", "Five-minute failure drill + submission", "Ready pack"),
            ],
            "outcomes": [
                "Diagnose a Python/venv activation failure using the smallest safe test.",
                "Diagnose a heuristic logic or evaluation bug without guessing randomly.",
                "Measure before/after metrics and explain what changed.",
                "Document failed attempts honestly.",
                "Explain when to stop and ask a supervisor.",
            ],
            "lab_title": "Controlled troubleshooting laboratory",
            "notebook": notebook,
            "script": script,
            "folder": folder,
            "verify_cmds": common_verify,
            "install_needed": False,
            "code": '''# day_04_lab.py — Intentionally brittle starter for troubleshooting practice
# Faults to find (do not skip):
# 1) Division risk / empty TEST handling
# 2) Train/test leakage if you evaluate on the same list you tuned
# 3) Keyword set that systematically misses a class

from typing import List, Tuple

SAMPLES: List[Tuple[str, bool]] = [
    ("server down", True),
    ("reset password", False),
    ("urgent payment failed", True),
    ("export csv help", False),
]

KEYWORDS = {"down", "urgent"}  # incomplete on purpose for the lab story

def predict(text: str) -> bool:
    return any(k in text.lower() for k in KEYWORDS)

def precision_recall(rows):
    tp = fp = fn = 0
    for text, truth in rows:
        pred = predict(text)
        if pred and truth: tp += 1
        elif pred and not truth: fp += 1
        elif (not pred) and truth: fn += 1
    precision = tp / (tp + fp)  # may ZeroDivisionError — catch and explain
    recall = tp / (tp + fn)
    return precision, recall

if __name__ == "__main__":
    try:
        p, r = precision_recall([])  # Fault A trigger if left empty
        print(p, r)
    except ZeroDivisionError as e:
        print("Caught metric failure:", type(e).__name__)
        print("Fix: guard denominators; never report fake metrics.")
    p, r = precision_recall(SAMPLES)
    print("Current precision/recall on SAMPLES:", round(p, 3), round(r, 3))
    print("Next: improve KEYWORDS OR document why a miss is acceptable; rerun; keep logs.")
''',
            "expected": "You capture the error, apply a safe fix, rerun, and show before/after evidence.",
            "fault_for_day4": True,
        }

    # Document and improve
    return {
        "one_liner": "Rewrite the Day 2–4 procedure so another intern can reproduce it, then improve one weak control (naming, validation, safety note, or metric honesty).",
        "safety": "Do not overwrite earlier evidence folders. Create Day_05 folders and reference prior days.",
        "schedule": [
            ("0:00–0:40", "Audit Day 2–4 folders for gaps", "Gap list"),
            ("0:40–2:10", "Rewrite README + procedure for reproducibility", "README.md"),
            ("2:10–2:25", "Break", "—"),
            ("2:25–4:25", "Implement one improvement + before/after proof", "Improvement evidence"),
            ("4:25–4:40", "Break", "—"),
            ("4:40–6:40", "Report + packaging", "Submission pack"),
            ("6:40–8:00", "Peer-style checklist + supervisor rehearsal", "Ready pack"),
        ],
        "outcomes": [
            "Document a reproducible environment and procedure.",
            "Improve one weak point with measurable before/after evidence.",
            "Organize filenames to the Beyond standard.",
            "Verify another person could follow your README without you present.",
            "Update limitations and responsible-AI controls after the improvement.",
        ],
        "lab_title": "Documentation and improvement laboratory",
        "notebook": notebook,
        "script": script,
        "folder": folder,
        "verify_cmds": common_verify,
        "install_needed": False,
        "code": '''# day_05_lab.py — Documentation helper: freeze versions + run a sanity metric
import platform
import sys

def main():
    print("python", sys.version.replace("\\n", " "))
    print("executable", sys.executable)
    print("platform", platform.platform())
    # Sanity: prove the heuristic module pattern still imports if you kept day_02_lab nearby
    print("OK: environment snapshot captured for README/report.")

if __name__ == "__main__":
    main()
''',
        "expected": "Environment snapshot captured; README explains exact steps; before/after improvement evidenced.",
        "fault_for_day4": None,
    }


def build_handbook(day_num: int) -> Path:
    day = load_day(day_num)
    if day_num == 1:
        raise SystemExit("Do not regenerate Day 1.")
    spec = mode_lab_spec(day)
    doc = new_doc_from_template()
    folder = spec["folder"]
    topic = day["topic"]
    mode = day["mode"]

    # --- 1 Title page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BEYOND ARTIFICIAL INTELLIGENCE INTERNSHIP")
    set_run(r, bold=True, size=14, color=BLUE)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"DAY {day_num} STUDENT HANDBOOK")
    set_run(r, bold=True, size=22, color=BLUE)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(topic)
    set_run(r, bold=True, size=14, color=DARK)

    add_p(doc, f"{mode}. Step-by-step laboratory manual with verification, troubleshooting, evidence, and report guidance.", size=11)

    add_table(
        doc,
        ["DAY MODE", "LEVEL", "DURATION", "CORE TOOLS"],
        [[mode, day.get("difficulty") or "foundation", f"{day.get('hours') or 8} hours", day.get("tools") or "Python, Jupyter, VS Code"]],
    )

    add_callout(doc, "note", f"Day {day_num} objective", spec["one_liner"])
    add_callout(doc, "safety", "Non-negotiable rule", spec["safety"])
    add_callout(
        doc,
        "warning",
        "Continuity",
        "Day 1 safety, evidence, and documentation rules remain in force. Do not reinstall the full toolchain unless verification fails and the supervisor approves.",
    )

    # --- 2 How to use ---
    doc.add_heading("How to use this handbook", level=1)
    add_p(
        doc,
        "This is a follow-along guide. Complete the actions in order, capture evidence at the marked checkpoints, "
        "and write your report while the work is still fresh. Do not skip verification. Activities that change shared "
        "accounts, networks, or production systems require supervisor approval—stop and ask.",
    )
    add_bullets(
        doc,
        [
            "Order: authorization → environment verify → lab steps → independent verification → evidence → report → checklists.",
            "Capture evidence when the checkpoint says so—not only at the end of the day.",
            f"Submit: {(day.get('submission') or 'evidence pack + technical report')}.",
            "Ask the supervisor before installing new system-wide software, using paid APIs, or touching non-lab systems.",
        ],
    )

    # --- 3 Schedule ---
    doc.add_heading("Suggested daily schedule", level=1)
    add_table(doc, ["TIME", "ACTIVITY", "OUTPUT"], spec["schedule"])

    # --- 4 Outcomes ---
    doc.add_heading("Learning outcomes", level=1)
    add_p(doc, f"Curriculum objective (source of truth): {day.get('objective')}", size=10)
    add_bullets(doc, spec["outcomes"])

    # --- 5 Key concepts ---
    doc.add_heading("Key concepts", level=1)
    doc.add_heading("What this topic means", level=2)
    add_p(
        doc,
        "AI orientation means knowing what kinds of AI systems exist, what problem you are solving, what data you may use, "
        "and how you will evaluate success. Responsible use means applying privacy, security, fairness, transparency, "
        "human oversight, and honest documentation controls while you work.",
    )
    add_callout(
        doc,
        "warning",
        "Writing correction",
        'Do not treat connecting words such as “and” as technical terms. Define “AI orientation,” “responsible use,” and the combined professional practice.',
    )
    add_table(
        doc,
        ["TERM", "PLAIN-LANGUAGE MEANING", "TODAY'S ACTION", "COMMON FAILURE", "CONTROL"],
        [
            ["Problem framing", "Who needs what, why, constraints, success", "Write a 6-row frame table", "Solving the wrong problem", "Supervisor review of frame before build"],
            ["Heuristic", "Explicit human-written rules", "Implement keyword/rule baseline", "Hidden brittle rules", "Keep rules visible in code + README"],
            ["Evaluation", "Measure on held-out examples", "Report TEST metrics", "Testing on training data", "Separate TEST list; rerun after restart"],
            ["Privacy", "Minimize exposure of personal data", "Use synthetic tickets only", "Pasting real messages into tools", "Data classification checklist"],
            ["Human oversight", "A person remains accountable", "You verify AI/tool claims", "Blind trust in model output", "AI assistance disclosure + verification"],
        ],
    )

    # --- 6 Prerequisites ---
    doc.add_heading("Prerequisites and authorization check", level=1)
    add_bullets(
        doc,
        [
            "Required prior learning: Day 1 orientation lab completed (Python, venv, VS Code, Jupyter, evidence habits).",
            "Authorized workstation / Beyond lab / VM only.",
            "Administrator permission only if your supervisor already approved installs.",
            "Disk: at least 5 GB free for the shared environment; more if using local models.",
            "No client/PII datasets. Synthetic data only unless formally authorized.",
            "Rollback: keep Day 1–previous day folders intact; create a new Day folder today.",
        ],
    )
    add_callout(doc, "safety", "Authorization", "If any checkbox is unclear, stop and ask the supervisor. Do not bypass security controls.")

    # --- 7 Download/install ---
    doc.add_heading("Download and installation guide", level=1)
    if not spec["install_needed"]:
        add_p(
            doc,
            "Full installation was covered in Day 1. Today you verify the existing toolchain. "
            "Reinstall only if verification fails and the supervisor approves.",
        )
        add_p(doc, "Official sources (if reinstall is approved):", bold=True)
        add_bullets(
            doc,
            [
                "Python: https://www.python.org/downloads/ (Python Software Foundation)",
                "VS Code: https://code.visualstudio.com/ (Microsoft)",
                "JupyterLab via pip inside your project venv (Project Jupyter)",
            ],
        )
        add_p(doc, "Verification commands (Windows: prefer py -m … / macOS-Linux: python3 -m …):", bold=True)
        add_code(doc, "\n".join(spec["verify_cmds"]))
        add_table(
            doc,
            ["PLATFORM", "ACTIVATE VENV", "EXPECTED"],
            [
                ["Windows PowerShell", r".venv\Scripts\Activate.ps1", "Prompt shows (.venv)"],
                ["Windows CMD", r".venv\Scripts\activate.bat", "Prompt shows (.venv)"],
                ["macOS / Linux", "source .venv/bin/activate", "Prompt shows (.venv)"],
            ],
        )
    add_callout(doc, "safety", "Download safety", "Use official websites only. Avoid software aggregators, cracks, and random pasted installers.")

    # --- 8 Project organization ---
    doc.add_heading("Project and file organization", level=1)
    add_code(
        doc,
        f"""Beyond_AI_Internship/
└── {folder}/
    ├── evidence/
    ├── notebooks/
    ├── src/
    ├── data/
    ├── output/
    ├── report/
    ├── requirements.txt
    └── README.md""",
    )
    add_p(doc, "Windows PowerShell:", bold=True)
    add_code(
        doc,
        f'''cd "$HOME\\Documents"
mkdir Beyond_AI_Internship\\{folder}\\evidence, Beyond_AI_Internship\\{folder}\\notebooks, Beyond_AI_Internship\\{folder}\\src, Beyond_AI_Internship\\{folder}\\data, Beyond_AI_Internship\\{folder}\\output, Beyond_AI_Internship\\{folder}\\report -Force
cd Beyond_AI_Internship\\{folder}''',
    )
    add_p(doc, "macOS / Linux:", bold=True)
    add_code(
        doc,
        f"""cd ~/Documents
mkdir -p Beyond_AI_Internship/{folder}/{{evidence,notebooks,src,data,output,report}}
cd Beyond_AI_Internship/{folder}""",
    )
    add_callout(doc, "warning", "Secrets", "Never place passwords, API keys, tokens, or personal records in the submission folder.")

    # --- 9 Environment ---
    doc.add_heading("Environment setup", level=1)
    add_numbered(
        doc,
        [
            "Open the Day folder in VS Code (File → Open Folder).",
            "Activate the existing .venv from Day 1 (or recreate only if approved).",
            "Select the .venv interpreter (Python: Select Interpreter).",
            "Confirm packages with python -m pip check and refresh requirements.txt if the supervisor asks.",
            "Do not put real keys in notebooks. If a lab key is approved later, use API_KEY=YOUR_APPROVED_LAB_KEY in a gitignored .env.",
        ],
    )
    add_code(
        doc,
        """python -m pip freeze > requirements.txt
python -m pip check""",
    )

    # --- 10 Step-by-step lab ---
    doc.add_heading("Step-by-step practical lab", level=1)
    add_p(doc, spec["lab_title"], bold=True, size=12, color=BLUE)
    add_p(doc, f"Create `src/{spec['script']}` with the following beginner-friendly starter (customize where instructed):")
    add_code(doc, spec["code"])

    doc.add_heading("Lab steps", level=2)
    steps = [
        (
            "Create folders and README",
            "Use the commands above",
            "Folder tree exists",
            "Proves organization discipline",
            f"day-{day_num:02d}_folder_tree.png",
            "Permission denied",
            "Use Documents path or ask IT—do not disable security",
        ),
        (
            "Activate venv and verify versions",
            "Activation command + python --version",
            "(.venv) prompt + version text",
            "Proves correct interpreter",
            f"day-{day_num:02d}_python_version.png",
            "Wrong Python",
            "Reselect interpreter; do not sudo-pip into system Python",
        ),
        (
            "Run the lab script",
            f"python src/{spec['script']}",
            spec["expected"],
            "Proves executable practical work",
            f"day-{day_num:02d}_script_output.png",
            "Module/import errors",
            "Activate venv; pip check; fix one issue at a time",
        ),
        (
            "Record primary metrics / outcomes",
            "Copy TEST metrics into report table",
            "Metrics table completed",
            "Proves evaluation literacy",
            f"day-{day_num:02d}_metrics.png",
            "Only train metrics reported",
            "Re-run on held-out TEST; label clearly",
        ),
        (
            "Independent verification",
            "Restart terminal, reactivate venv, rerun script",
            "Second run matches first within rounding",
            "Proves reproducibility",
            f"day-{day_num:02d}_rerun.png",
            "Different results",
            "Check randomness, working directory, interpreter path",
        ),
    ]
    for i, (action, cmd, exp, proves, evid, fail, fix) in enumerate(steps, 1):
        doc.add_heading(f"Step {i}: {action}", level=2)
        add_table(
            doc,
            ["FIELD", "DETAIL"],
            [
                ["Action", action],
                ["Command / work", cmd],
                ["Expected result", exp],
                ["What it proves", proves],
                ["Evidence to capture", evid],
                ["Likely failure", fail],
                ["Safe corrective action", fix],
            ],
        )

    # --- 11 Responsible AI ---
    doc.add_heading("Responsible-AI checkpoint", level=1)
    add_table(
        doc,
        ["RISK", "POSSIBLE HARM", "PREVENTION", "DETECTION", "APPROVER"],
        [
            ["Privacy leakage", "Exposure of personal/client data", "Synthetic data only; no public paste", "Screenshot privacy review", "Student + supervisor"],
            ["Overconfident AI text", "Wrong operational decision", "Treat AI as draft; verify claims", "Manual check of facts", "Student"],
            ["Bias in rules/keywords", "Uneven treatment of groups/topics", "Review keyword list for skew", "Error analysis by category", "Student + supervisor if live use"],
            ["Unapproved automation", "Spam/external side effects", "No outbound bots/APIs without approval", "Code review for network calls", "Supervisor"],
            ["Irreproducible work", "Cannot audit internship evidence", "Freeze requirements; rerun proof", "Second-run evidence", "Student"],
        ],
    )

    # --- 12 Evaluation ---
    doc.add_heading("Evaluation and metrics", level=1)
    add_p(doc, "Use only metrics that fit today's heuristic/classifier exercise.")
    add_table(
        doc,
        ["METRIC", "QUESTION IT ANSWERS", "CAUTION"],
        [
            ["Accuracy", "Overall fraction correct", "Misleading on imbalanced labels"],
            ["Precision", "Of predicted urgent, how many truly urgent?", "Can look good while missing many urgents"],
            ["Recall", "Of true urgents, how many found?", "Can look good with many false alarms"],
            ["F1", "Balance of precision and recall", "Still needs human context for cost of errors"],
        ],
    )
    add_p(doc, "Baseline: a naive rule (for example, never urgent) compared against your heuristic. Success: you can explain errors, not merely quote a high score.")

    # --- 13 Independent verification ---
    doc.add_heading("Independent verification", level=1)
    add_bullets(
        doc,
        [
            "Restart the terminal, reactivate .venv, and rerun the script.",
            "Manually classify 4 TEST rows on paper and compare to the program.",
            "Confirm requirements.txt matches the environment used for the rerun.",
            "A screenshot alone is evidence of activity, not proof of correctness—pair it with the rerun and manual check.",
        ],
    )

    # --- 14 Troubleshooting ---
    doc.add_heading("Troubleshooting guide", level=1)
    add_table(
        doc,
        ["SYMPTOM", "LIKELY CAUSE", "SAFE FIRST ACTION", "VERIFY AFTER", "ESCALATE IF"],
        [
            ["venv activate blocked (PS)", "Execution policy", "Use CMD activate or ask IT—do not weaken org policy", "(.venv) shows", "Policy forbids all script activation"],
            ["ModuleNotFoundError", "Wrong interpreter", "Select .venv interpreter; pip show package", "Import works", "Package cannot install on locked image"],
            ["Metrics ZeroDivisionError", "Empty pred/positive set", "Guard denominators; check labels", "Script prints metrics", "Labels/data corrupt"],
            ["Results differ after rerun", "Different cwd/interpreter", "Print sys.executable and cwd", "Two runs match", "Nondeterminism you cannot control"],
            ["Jupyter kernel missing", "Kernel not using .venv", "Python: Select Interpreter; restart Jupyter", "Notebook runs", "VS Code extensions blocked"],
        ],
    )
    add_callout(
        doc,
        "warning",
        "Five-minute failure procedure",
        "Before a demo: (1) correct folder, (2) (.venv) active, (3) python --version, (4) rerun script, (5) open evidence filenames, (6) close windows that show secrets/notifications.",
    )
    add_p(doc, "Students must not: hide errors, fabricate results, disable security controls, modify production equipment, delete failure evidence, download unapproved software, or expose a local server to the public internet without approval.")

    # --- 15 Evidence ---
    doc.add_heading("Evidence checkpoints", level=1)
    add_table(
        doc,
        ["FILE", "MUST SHOW", "WHY", "PRIVACY CHECK", "FORMAT"],
        [
            [f"day-{day_num:02d}_python_version.png", "Python version + venv path", "Environment authenticity", "No home directory secrets", "PNG/JPEG"],
            [f"day-{day_num:02d}_script_output.png", "Script output / metrics", "Technical execution", "No API keys", "PNG/JPEG"],
            [f"day-{day_num:02d}_rerun.png", "Second successful run", "Verification", "No notifications", "PNG/JPEG"],
            [f"src/{spec['script']}", "Commented beginner code", "Reproducibility", "No secrets in source", "PY"],
            [f"report/day-{day_num:02d}_technical_report.md", "Full report sections", "Professional communication", "No PII", "MD/PDF"],
        ],
    )

    # --- 16 Report ---
    doc.add_heading("Professional technical report guide", level=1)
    add_p(doc, "Required structure:", bold=True)
    add_numbered(
        doc,
        [
            "Title and metadata (name, day, date, program)",
            "Objective",
            "Scope and authorization",
            "Concept summary",
            "Environment and versions",
            "Data or input description (synthetic)",
            "Procedure",
            "Results",
            "Evaluation metrics",
            "Evidence references",
            "Safety, ethics, and risk controls",
            "Independent verification",
            "Troubleshooting",
            "Limitations",
            "Conclusion and next step",
            "AI assistance disclosure",
            "References",
        ],
    )
    add_callout(
        doc,
        "note",
        "Example AI disclosure",
        "I used [approved tool/model] to assist with [specific task]. I did not provide personal or client data. "
        "I independently verified [claims/code/results] using [method/source]. I corrected or rejected [brief description], "
        "and I remain responsible for the submitted work.",
    )

    # --- 17 Submission checklists ---
    doc.add_heading("Submission checklists", level=1)
    doc.add_heading("A. Safety and ethics checklist", level=2)
    add_bullets(
        doc,
        [
            "Only synthetic / approved data used",
            "No secrets in screenshots, code, or prompts",
            "No production or shared-system changes",
            "AI assistance disclosed",
            "Time entries truthful",
        ],
    )
    doc.add_heading("B. Technical completion checklist", level=2)
    add_bullets(
        doc,
        [
            f"Day folder {folder} created",
            "venv verified and script/notebook executed",
            "TEST/primary metrics recorded",
            "Independent verification completed",
            "Troubleshooting notes include at least one real issue or a supervised simulated fault",
        ],
    )
    doc.add_heading("C. Evidence and report checklist", level=2)
    add_bullets(
        doc,
        [
            "Required screenshots present with correct filenames",
            "Source file submitted",
            "Report includes all required sections",
            "Filenames follow day-XX_topic_short-description.ext",
            f"Submission matches: {(day.get('submission') or 'evidence + report')}",
        ],
    )

    # --- 18 Self-check ---
    doc.add_heading("Self-check questions", level=1)
    add_numbered(
        doc,
        [
            "Explain AI orientation and responsible use in under two minutes without notes.",
            "Why is a heuristic acceptable for today's lab, and when would you consider a learned model?",
            "What is the difference between evidence and proof in your submission?",
            "Name one privacy failure mode and the control you applied.",
            "Which metric can look good while still missing urgent cases?",
            "What is your first safe action if venv activation fails?",
            "Why is this work not production-ready?",
            "What would you ask the supervisor before using a public LLM API?",
        ],
    )

    # --- 19 Two-minute explanation ---
    doc.add_heading("Two-minute supervisor explanation", level=1)
    add_p(
        doc,
        f"Today I completed Day {day_num} ({mode}) on {topic}. The problem was to practise responsible AI with a transparent "
        f"decision aid on synthetic data. My method was a documented heuristic plus held-out evaluation. The result was a "
        f"reproducible script/notebook with metrics and evidence. I verified by restarting the environment and manually "
        f"checking sample labels. The main risk is privacy leakage or overconfident automation; the control is synthetic data, "
        f"no outbound actions, and human review. The limitation is that keyword rules are brittle and not production-ready. "
        f"Next I will continue the topic cycle and strengthen documentation and evaluation discipline.",
    )

    # --- 20 Rubric ---
    doc.add_heading("Supervisor assessment rubric", level=1)
    add_table(
        doc,
        ["AREA", "PASS EVIDENCE", "WEIGHT"],
        [
            ["Safe and authorized setup", "Checklist + lab-only tools", "15%"],
            ["Technical execution", "Working script/notebook + outputs", "25%"],
            ["Concept understanding", "Correct definitions in report/oral", "15%"],
            ["Evaluation and verification", "TEST metrics + rerun proof", "15%"],
            ["Responsible-AI application", "Risk table + disclosure", "10%"],
            ["Evidence quality", "Clear filenames, no secrets", "10%"],
            ["Professional reporting", "Complete report structure", "5%"],
            ["Honest troubleshooting", "Real fault notes / no fabrication", "5%"],
        ],
    )
    add_p(doc, f"Pass mark from curriculum: {day.get('pass_mark', 60)}%.")

    # --- 21 References ---
    doc.add_heading("Official references", level=1)
    add_p(doc, f"Access check date for link validity: {date.today().isoformat()}. Prefer first-party docs; do not invent titles or URLs.")
    add_bullets(
        doc,
        [
            "Python Software Foundation — https://docs.python.org/3/",
            "Python downloads — https://www.python.org/downloads/",
            "Visual Studio Code docs — https://code.visualstudio.com/docs",
            "Jupyter documentation — https://docs.jupyter.org/",
            "NIST AI Risk Management Framework (AI RMF 1.0) — https://www.nist.gov/itl/ai-risk-management-framework",
            "NIST Privacy Framework — https://www.nist.gov/privacy-framework",
            "OWASP Top Ten / LLM risks (as applicable) — https://owasp.org/",
        ],
    )
    add_callout(
        doc,
        "note",
        "Instructor Review Required",
        "If your site bans local model downloads or public LLM use, follow local Beyond policy over any optional generative exercise.",
    )

    out = HERE / f"Beyond_AI_Internship_Day_{day_num:03d}_Student_Handbook.docx"
    doc.save(out)
    return out


def update_index_status(days_done):
    """Best-effort: append changelog; full index rebuild via build_phase1_audit for structure."""
    lines = []
    if CHANGELOG.exists():
        lines = CHANGELOG.read_text().splitlines()
    else:
        lines = ["# AI Handbook changelog", ""]
    lines.append(f"## {date.today().isoformat()}")
    for d in days_done:
        lines.append(f"- Generated Beyond_AI_Internship_Day_{d:03d}_Student_Handbook.docx (pilot batch)")
    lines.append("")
    CHANGELOG.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--days", default="2-5", help="e.g. 2-5 or 2,3,4,5")
    args = ap.parse_args()
    if "-" in args.days and "," not in args.days:
        a, b = args.days.split("-", 1)
        day_list = list(range(int(a), int(b) + 1))
    else:
        day_list = [int(x) for x in args.days.split(",")]

    done = []
    for d in day_list:
        if d == 1:
            print("Skipping Day 1 (canonical).")
            continue
        path = build_handbook(d)
        print("Wrote", path.name)
        done.append(d)
    update_index_status(done)
    print("Updated", CHANGELOG.name)


if __name__ == "__main__":
    main()
