#!/usr/bin/env python3
"""Phase 1: AI curriculum audit, dependency map, and master index DOCX."""

from __future__ import annotations

import json
from collections import OrderedDict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SEED = HERE.parents[1] / "beyond_180_day_curriculum_seed.json"
COMPACT = HERE / "_ai_days_compact.json"
AUDIT_MD = HERE / "Beyond_AI_Internship_Phase1_Audit_Report.md"
DEP_JSON = HERE / "curriculum_dependency_map.json"
INDEX_DOCX = HERE / "Beyond_AI_Internship_Handbook_Master_Index.docx"
DAY001 = HERE / "Beyond_AI_Internship_Day_001_Student_Handbook.docx"

PHASES = [
    ("Phase 1 — Orientation & tooling", 1, 24, "AI orientation, Python/VS Code/Jupyter, data structures, Git & reproducible notebooks"),
    ("Phase 2 — Classic AI foundations", 25, 48, "Problem framing, search/optimization, knowledge representation, rule-based systems"),
    ("Phase 3 — Data & math for AI", 49, 66, "Data preparation, statistics for AI, linear algebra intuition"),
    ("Phase 4 — Classical machine learning", 67, 102, "ML overview through feature engineering and model evaluation"),
    ("Phase 5 — Deep learning & modalities", 103, 132, "Neural nets, CV, NLP, speech/audio, recommendation systems"),
    ("Phase 6 — Generative AI, RAG & agents", 133, 156, "Generative AI, prompting, RAG, agents & tool use"),
    ("Phase 7 — Governance, deployment & capstone", 157, 180, "Bias/privacy/safety, deployment/monitoring, local-business prototype, capstone"),
]

MODE_ORDER = [
    "Orientation and setup",
    "Guided practical",
    "Independent build",
    "Troubleshoot",
    "Document and improve",
    "Assessment and reflection",
]


def load_days():
    if COMPACT.exists():
        return json.loads(COMPACT.read_text())
    data = json.loads(SEED.read_text())
    ai = next(x for x in data["programs"] if x["code"] == "ARTIFICIAL_INTELLIGENCE")
    rows = []
    for t in ai["tasks"]:
        title = t["title"]
        mode, topic = (title.split(":", 1) + [""])[:2] if ":" in title else ("?", title)
        rows.append(
            {
                "day": t["day_number"],
                "mode": mode.strip(),
                "topic": topic.strip(),
                "difficulty": t.get("difficulty"),
                "hours": t.get("estimated_hours"),
                "tools": t.get("tools"),
                "objective": t.get("objective"),
                "submission": t.get("submission"),
                "instructions": t.get("instructions"),
                "pass_mark": t.get("pass_mark"),
            }
        )
    return rows


def topic_modules(days):
    modules = OrderedDict()
    for d in days:
        modules.setdefault(d["topic"], []).append(d)
    return modules


def phase_for(day_num: int) -> str:
    for name, a, b, _ in PHASES:
        if a <= day_num <= b:
            return name
    return "Unassigned"


def build_dependency_map(days):
    modules = topic_modules(days)
    topics = list(modules.keys())
    map_out = {
        "program": "ARTIFICIAL_INTELLIGENCE",
        "schema_version": 1,
        "generated": str(date.today()),
        "mode_cycle": MODE_ORDER,
        "phases": [
            {"name": n, "day_start": a, "day_end": b, "summary": s} for n, a, b, s in PHASES
        ],
        "topics": [],
        "day_prerequisites": {},
    }

    carried_tools = "Python + venv + VS Code + Jupyter (from Day 1)"
    for i, topic in enumerate(topics):
        topic_days = modules[topic]
        start = topic_days[0]["day"]
        end = topic_days[-1]["day"]
        prev = topics[i - 1] if i else None
        entry = {
            "index": i + 1,
            "topic": topic,
            "days": [d["day"] for d in topic_days],
            "day_start": start,
            "day_end": end,
            "difficulty": topic_days[0]["difficulty"],
            "tools": topic_days[0]["tools"],
            "depends_on_topics": ([prev] if prev else []) + (["Ai Orientation And Responsible Use"] if i > 0 else []),
            "carried_environment": carried_tools,
            "notes": [],
        }
        if i == 0:
            entry["notes"].append("Day 1 installs the shared toolchain; Days 2–6 deepen the same topic without reinstalling by default.")
        if "Git" in topic:
            entry["notes"].append("Introduces Git; later days assume commits and reproducibility habits.")
        if topic in ("Machine Learning Overview", "Classification Concepts"):
            entry["depends_on_topics"].extend(["Data Preparation", "Statistics For Ai"])
        if "Neural" in topic or "Computer Vision" in topic:
            entry["notes"].append("May need more RAM/disk; local GPU optional — Instructor Review if required.")
        if any(k in topic for k in ("Generative", "Prompt", "Retrieval", "Agents")):
            entry["notes"].append("External/public model APIs need supervisor approval; prefer local/open or approved lab keys.")
        if "Deployment" in topic or "Prototype" in topic or "Capstone" in topic:
            entry["notes"].append("No production deploy without written supervisor approval.")
        map_out["topics"].append(entry)

        for d in topic_days:
            prereq_days = []
            if d["day"] > 1:
                prereq_days.append(1)  # safety + toolchain
            # previous day in cycle
            if d["day"] > start:
                prereq_days.append(d["day"] - 1)
            # previous topic assessment day
            if i > 0 and d["mode"] == "Orientation and setup":
                prereq_days.append(modules[prev][-1]["day"])
            map_out["day_prerequisites"][str(d["day"])] = {
                "day": d["day"],
                "mode": d["mode"],
                "topic": topic,
                "phase": phase_for(d["day"]),
                "prerequisite_days": sorted(set(prereq_days)),
                "assumes": [
                    "Authorized Beyond lab/workstation",
                    carried_tools if d["day"] > 1 else "Fresh authorized workstation",
                    "Day 1 safety and evidence standards remain in force",
                ],
            }
    return map_out


def build_flags(days, modules):
    flags = {
        "duplicates_overlap": [
            {
                "item": "Days 1–6 share topic Ai Orientation And Responsible Use",
                "detail": "Day 1 is the install + baseline; Days 2–6 must add guided/independent/troubleshoot/document/assessment depth—not repeat full installs.",
                "severity": "info",
            },
            {
                "item": "Python Environment Setup (Days 7–12) overlaps Day 1 toolchain",
                "detail": "Treat Day 7 as environment hardening, package hygiene, and OS-specific repair—not a second full Python install guide.",
                "severity": "medium",
            },
            {
                "item": "Generative AI / Prompt Engineering / RAG / Agents (Days 133–156)",
                "detail": "Topics overlap conceptually; each 6-day block must add a distinct practical skill and preserve prior artifacts.",
                "severity": "medium",
            },
            {
                "item": "Machine Learning Overview vs later Classification/Regression/Clustering",
                "detail": "Overview should stay conceptual + tiny demo; later topics own full model labs.",
                "severity": "info",
            },
        ],
        "missing_prerequisites": [
            {
                "item": "No explicit Git before Day 19",
                "detail": "Acceptable if Days 1–18 keep evidence in folders only; introduce Git carefully on Day 19.",
                "severity": "low",
            },
            {
                "item": "Linear Algebra / Statistics before Neural Nets",
                "detail": "Present in curriculum order (Stats 55–60, LinAlg 61–66 before Neural 103+). Handbooks must reference those days.",
                "severity": "info",
            },
            {
                "item": "No dedicated SQL/database topic",
                "detail": "Data Preparation may need Assumption: tabular CSV/JSON only unless supervisor adds DB access.",
                "severity": "medium",
            },
        ],
        "unrealistic_scope": [
            {
                "item": "Speech And Audio Ai Basics (121–126)",
                "detail": "Needs microphone/audio files; may exceed 8h if models are large. Flag hardware + time.",
                "severity": "high",
            },
            {
                "item": "Computer Vision Basics (109–114)",
                "detail": "Dataset download and training can overrun 8h; use tiny synthetic/public subsets.",
                "severity": "high",
            },
            {
                "item": "Local-Business Ai Prototype + Capstone (169–180)",
                "detail": "True business prototype in 8h/day is ambitious; scope to Beyond-approved mini-prototype with clear MVP.",
                "severity": "high",
            },
            {
                "item": "Ai Agents And Tool Use (151–156)",
                "detail": "Tool-calling agents can touch external systems—sandbox only; supervisor approval required.",
                "severity": "high",
            },
        ],
        "special_requirements": [
            {"item": "Local/open models", "days": "1–180 (optional)", "need": "Disk/RAM; supervisor approval"},
            {"item": "Paid/public LLM APIs", "days": "133–156 especially", "need": "Approved lab key; never personal billing"},
            {"item": "Camera/mic", "days": "109–126", "need": "Optional; synthetic media preferred"},
            {"item": "Deployment target", "days": "163–168", "need": "Authorized sandbox only"},
        ],
        "safety_legal": [
            {"item": "No client/PII in prompts or datasets", "applies": "all days"},
            {"item": "No production scanning, scraping, or automated outbound actions", "applies": "agents, RAG, deployment, prototype"},
            {"item": "AI outputs are not professional advice without human review", "applies": "all generative days"},
            {"item": "Disclose AI assistance; keep failed attempts", "applies": "all days"},
        ],
    }
    return flags


def write_audit_md(days, dep, flags):
    modules = topic_modules(days)
    lines = [
        "# Beyond AI Internship — Phase 1 Audit Report",
        "",
        f"**Generated:** {date.today().isoformat()}",
        "**Program:** ARTIFICIAL_INTELLIGENCE (180 days)",
        "**Canonical Day 1:** `Beyond_AI_Internship_Day_001_Student_Handbook.docx` (do not rewrite)",
        "**Follow-on scope:** After AI handbooks complete, apply the same format to the other 10 internship programs.",
        "",
        "## 1. Executive summary",
        "",
        "The AI curriculum is a clean **30-topic × 6-mode** ladder. Day 1 already provides professional-depth install and responsible-AI baseline. "
        "Days 2–180 must become full student handbooks (same depth as Day 1) with mode-specific labs, without silently changing learning objectives.",
        "",
        f"- Topics: **{len(modules)}**",
        f"- Days audited: **{len(days)}**",
        f"- Day 001 handbook present: **{DAY001.exists()}**",
        "",
        "## 2. Program-phase breakdown",
        "",
    ]
    for name, a, b, summary in PHASES:
        lines.append(f"### {name} (Days {a}–{b})")
        lines.append("")
        lines.append(summary)
        lines.append("")
        topic_set = OrderedDict()
        for d in days:
            if a <= d["day"] <= b:
                topic_set.setdefault(d["topic"], True)
        for t in topic_set:
            lines.append(f"- {t}")
        lines.append("")

    lines.extend(
        [
            "## 3. Mode cycle (every topic)",
            "",
            "| Step | Mode | Intent for handbook |",
            "|---|---|---|",
            "| 1 | Orientation and setup | Verify environment; introduce topic concepts; light setup lab |",
            "| 2 | Guided practical | Full step-by-step lab with expected outputs |",
            "| 3 | Independent build | Student-owned artifact; less scaffolding |",
            "| 4 | Troubleshoot | Controlled fault; hypothesis → test → fix |",
            "| 5 | Document and improve | Reproducibility, naming, before/after improvement |",
            "| 6 | Assessment and reflection | End-to-end demo + honest reflection + rubric |",
            "",
            "## 4. Topic dependency map (summary)",
            "",
            "| # | Topic | Days | Difficulty | Key dependencies |",
            "|---|---|---|---|---|",
        ]
    )
    for t in dep["topics"]:
        deps = ", ".join(t["depends_on_topics"][:3]) or "—"
        lines.append(
            f"| {t['index']} | {t['topic']} | {t['day_start']}–{t['day_end']} | {t['difficulty']} | {deps} |"
        )

    lines.extend(["", "## 5. Duplicates and overlap", ""])
    for item in flags["duplicates_overlap"]:
        lines.append(f"- **{item['item']}** ({item['severity']}): {item['detail']}")

    lines.extend(["", "## 6. Missing prerequisites", ""])
    for item in flags["missing_prerequisites"]:
        lines.append(f"- **{item['item']}** ({item['severity']}): {item['detail']}")

    lines.extend(["", "## 7. Unrealistic duration or scope", ""])
    for item in flags["unrealistic_scope"]:
        lines.append(f"- **{item['item']}** ({item['severity']}): {item['detail']}")

    lines.extend(["", "## 8. Special hardware / services / credentials", ""])
    for item in flags["special_requirements"]:
        lines.append(f"- **{item['item']}** — days {item['days']}: {item['need']}")

    lines.extend(["", "## 9. Safety / legal sensitivity", ""])
    for item in flags["safety_legal"]:
        lines.append(f"- **{item['item']}** — applies: {item['applies']}")

    lines.extend(
        [
            "",
            "## 10. Handbook production recommendation",
            "",
            "1. Keep Day 001 unchanged.",
            "2. Generate Days 002–005 as pilot (same topic, modes Guided → Document).",
            "3. Continue in batches of five through Day 180.",
            "4. After AI completes, repeat Phase 1+pilot for each remaining program.",
            "5. Do not invent production credentials, client data, or unpaid cloud spend.",
            "",
            "## 11. Approval gate",
            "",
            "Approve this audit (or list corrections) before treating Days 006+ as authorized for mass production. "
            "Pilot Days 002–005 may proceed immediately as the format validation batch.",
            "",
            "## 12. Other programs (queued after AI)",
            "",
            "NETWORKING, CYBER_SECURITY, CLOUD_COMPUTING, MACHINE_LEARNING, DATA_SCIENCE, "
            "SOFTWARE_DEVELOPMENT, LIVE_SOUND_ENGINEERING, LIGHTING_ENGINEERING, SCREENS_VIDEO, INTERCOM.",
            "",
        ]
    )
    AUDIT_MD.write_text("\n".join(lines), encoding="utf-8")


def shade_header_row(row):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0B3F90")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(9)


def write_master_index(days, dep):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BEYOND ARTIFICIAL INTELLIGENCE INTERNSHIP")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0B, 0x3F, 0x90)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Handbook Master Index")
    r.bold = True
    r.font.size = Pt(20)

    meta = doc.add_paragraph(
        f"Generated {date.today().isoformat()} · 180 days · Day 001 approved (external) · "
        "Days 002–180 status tracked below · Other programs queued after AI."
    )
    meta.runs[0].font.size = Pt(10)

    doc.add_heading("Program phases", level=1)
    pt = doc.add_table(rows=1, cols=4)
    pt.style = "Table Grid"
    hdr = pt.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Days"
    hdr[2].text = "Topics focus"
    hdr[3].text = "Handbook status"
    shade_header_row(pt.rows[0])
    for name, a, b, summary in PHASES:
        row = pt.add_row().cells
        row[0].text = name
        row[1].text = f"{a}–{b}"
        row[2].text = summary
        if b == 1:
            row[3].text = "Day 001 done"
        elif a <= 5:
            row[3].text = "Pilot in progress (002–005)"
        else:
            row[3].text = "Queued"

    doc.add_heading("Day-by-day index", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = [
        "Day",
        "Topic",
        "Mode / phase",
        "Difficulty",
        "Tools",
        "Main practical output",
        "Prerequisites",
        "Submission",
        "Status / filename",
    ]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    shade_header_row(table.rows[0])

    for d in days:
        pr = dep["day_prerequisites"][str(d["day"])]
        fname = f"Beyond_AI_Internship_Day_{d['day']:03d}_Student_Handbook.docx"
        if d["day"] == 1:
            status = f"Approved · {fname}"
        else:
            status = f"Not started · {fname}"
        cells = table.add_row().cells
        cells[0].text = str(d["day"])
        cells[1].text = d["topic"]
        cells[2].text = f"{d['mode']} · {phase_for(d['day']).split('—')[0].strip()}"
        cells[3].text = str(d["difficulty"] or "")
        cells[4].text = (d["tools"] or "")[:80]
        cells[5].text = (d["submission"] or "")[:100]
        cells[6].text = ", ".join(str(x) for x in pr["prerequisite_days"]) or "—"
        cells[7].text = (d["submission"] or "")[:80]
        cells[8].text = status
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_heading("Production notes", level=1)
    doc.add_paragraph(
        "Preserve learning objectives from the curriculum seed. When a day is vague, add an Assumption "
        "or Instructor Review Required callout. Never invent production credentials or client data."
    )
    doc.add_paragraph(
        "After the AI series is complete, create parallel folders under handbooks/ for each remaining program "
        "and repeat Phase 1 + pilot before full generation."
    )
    doc.save(INDEX_DOCX)


def main():
    days = load_days()
    dep = build_dependency_map(days)
    flags = build_flags(days, topic_modules(days))
    DEP_JSON.write_text(json.dumps(dep, indent=2), encoding="utf-8")
    write_audit_md(days, dep, flags)
    write_master_index(days, dep)
    print(f"Wrote {AUDIT_MD.name}")
    print(f"Wrote {DEP_JSON.name}")
    print(f"Wrote {INDEX_DOCX.name}")


if __name__ == "__main__":
    main()
